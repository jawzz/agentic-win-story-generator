"""
Win Story Generator - Flask Backend v2
Serves the frontend, handles AI extraction, and generates a single-slide
agentic use case PPTX.

Required env var: ANTHROPIC_API_KEY
"""
import json
import os
import traceback
from pathlib import Path

import requests as http_requests
from flask import Flask, request, send_from_directory, Response, jsonify
from flask_cors import CORS

from generate_pptx import build_pptx

app = Flask(__name__, static_folder='static')
CORS(app)


EXTRACTION_PROMPT = """You are a UiPath customer success writer. A rep has shared raw notes about an agentic automation win. Transform these notes into a polished, structured story for a single-slide agent use case template.

HARD CAPABILITY MAPPING RULES (apply before everything else):
- DOCUMENT UNDERSTANDING vs IXP: There are two valid product names for document/communication processing. Decide which to use based on what the rep wrote:
  * Use "DU" (the older Document Understanding product) ONLY when the notes EXPLICITLY say "Document Understanding", "Doc Understanding", or "DU" (older deployments often pre-date IXP and the rep will name DU directly).
  * Use "IXP" in ALL OTHER document/communication/form processing cases — including when the notes mention "Communications Mining", "intelligent document processing", "IDP", "OCR", or describe the work without naming a product.
  * Treat DU and IXP as interchangeable for layout purposes — they share the same green icon and slide treatment. The label is what differs.
- DOCUMENT/COMMUNICATION/FORM PROCESSING IS MANDATORY (mapped to either DU or IXP per the rule above) whenever ANY step pulls, reads, extracts, classifies, parses, ingests, or analyzes information from:
  * A document type: PDF, invoice, claim, contract, statement, remittance, RFQ/RFP, purchase order/PO, ACK/ASN/EDI, receipt, application, letter, memo, packing slip, bill of lading, fax, scanned image, screenshot, photo of a document.
  * A communication: email, email thread, inbox triage, voicemail, call transcript, chat message, SMS, ticket text, support case body, social message, customer message.
  * A form: web form, intake form, application form, questionnaire, survey, submission, claim form, onboarding form, KYC form, attachment.
  * Anything described as "extract fields from", "read X", "parse X", "pull data from X", "ingest X", "intake X", "classify X" where X is one of the above artifacts.
  When this happens you MUST:
  1. Add "DU" OR "IXP" to capabilities (per rule above). Never list "Doc Understanding" or "Communications Mining" as their own separate items — collapse them into DU or IXP.
  2. Use role "DU" or "IXP" (matching the capability) for that step — never "BOT" and never "AGENT". Description should name what's being processed, e.g. "Extract claim fields via DU", "Triage emails with IXP", "Read invoice line items via DU".
- MAESTRO IS MANDATORY whenever orchestration appears anywhere in the notes. Treat any of the following as an orchestration signal:
  * Explicit mentions: "Maestro", "orchestration", "orchestrate", "orchestrator", "process orchestrator", "workflow engine".
  * Coordination patterns: routing, dispatch, handoff between agents/bots/humans, queue management, SLA management, escalation, prioritization across queues, workload registration, workflow registration, process controller.
  * Multi-step solutions where work moves across more than one actor type (agents, bots, humans, IXP). Any solution with 3+ steps spanning multiple roles MUST include Maestro by default — orchestration is implicit.
- A "--- SUPPORTING DOCS ---" section (PDD, process map, transcript) attached to the notes is the source of truth for the steps. Map the solution flow and step roles closely to whatever the process map describes. If the process map has 12 low-level steps, abstract up one level so the slide has 5-7 steps that capture the key transitions, but preserve the order and the role at each transition.

TRANSPARENCY RULES FOR NUMBERS (follow carefully):
Every numerical metric (in problem_stats or outcomes) MUST be tagged with a "source" indicating where it came from:
- "stated"     = the exact number (or clear equivalent) appears in the notes. No derivation needed.
- "calculated" = you did simple arithmetic from numbers in the notes. Show the math in "note".
- "estimated"  = you inferred a plausible number from industry benchmarks or scope context. Label as estimate in "note".
- "qualitative"= no number given and none can be responsibly estimated. Use a word like "Reduced", "Faster", "Yes" as the value and leave source=qualitative.

RULES:
1. Every outcome and problem_stat item MUST have a "source" field.
2. Items with source = "calculated" or "estimated" MUST have a "note" field (max ~80 chars) explaining the derivation. Examples:
   - "From $500K/yr budget ÷ 12 months"
   - "Industry avg for claim triage (est.)"
   - "3,000 tickets × $120 handle cost"
3. Items with source = "stated" should still include a short "note" only if helpful (e.g. "per rep notes"). Usually the note can be empty.
4. Prefer "stated" whenever possible. Use "calculated" when basic arithmetic is obvious. Use "estimated" sparingly — and when you do, be honest about it in the note.
5. Never fabricate a specific number with source = "stated". That's a lie. If the rep didn't give you the number, mark it "estimated" or "calculated" and show your work, or use qualitative.
6. Do not invent capabilities, products, steps, agents, bots, or humans that aren't described in the notes. Anti-fabrication still applies to NON-numeric fields.
7. Encourage the rep to go find real numbers — every estimated/calculated outcome will show the derivation note in the UI so they can verify or replace it.

Rules (follow strictly):
- TITLE: SHORT and punchy. 3-6 words. Just the subject of the use case (e.g. "Medical claim adjudication", "Automating denials intake", "KYC refresh triage"). Do NOT include dollar amounts, percentages, or outcomes in the title - those go in the outcomes tiles. Default to the cleanest possible short title.
- SUBTITLE: one line, 8-14 words, summarizing the actor orchestration (e.g. "Agents reason over denials; bots execute; humans approve the edge cases.")
- PROBLEM DESCRIPTION: 2 concise sentences describing the pain. Professional tone. Max ~240 characters.
- PROBLEM STATS: up to 4 quantitative problem metrics (dollars, percentages, volumes, time). Each item: {value, label, source, note}. See transparency rules above. Prefer stated numbers. Empty array if you can't produce even one honestly.
- SOLUTION DESCRIPTION: 1-2 concise sentences describing how agents, bots, and humans resolve the problem. Max ~220 characters.
- CAPABILITIES: 3-6 UiPath product names/capabilities used. Canonical names (use these exactly):
    - "Agents" - AI reasoning, autonomous decisioning
    - "Maestro" - the orchestration layer. ALWAYS include Maestro when the solution involves coordinating agents, bots, or humans across a process, a workflow engine, process/workload orchestration, routing, handoffs, SLA management, or anything the notes call "workload registration", "workflow registration", "orchestration", or "process controller". Default to including Maestro on any multi-step agentic solution.
    - "IXP" - the umbrella term for intelligent document + communications processing. Default name. Use IXP for any modern document/communication/form processing.
    - "DU" - Document Understanding, the older product. Use ONLY when the rep notes explicitly mention "Document Understanding" or "DU" (older deployments). Same green slide treatment as IXP. Never list "Communications Mining" as separate — collapse into IXP.
    - "Unattended Robots" - deterministic RPA, system-to-system automation
    - "Attended Robots" - desktop assistant bots with humans
    - "Action Center" - human-in-the-loop approvals and reviews
    - "Data Service" - structured data storage
    - "API Integration" - integrations to external systems
    - "Test Suite", "Insights", "Apps" - only if clearly referenced
  If notes mention orchestration/routing/handoff/workflow engine: include Maestro. If notes mention any document/form/invoice/email extraction: include IXP (not Doc Understanding).
- STEPS: 3-9 discrete steps that are actually described in the notes. Each has a role:
  - "AGENT" = AI reasoning, classification, routing, decisions
  - "BOT"   = deterministic RPA (data entry, API calls, portal polling, system updates)
  - "HUMAN" = human in the loop (review, approve, sign-off)
  - "IXP"   = documents and communications processing — the umbrella for any step that parses PDFs, emails, forms, invoices, claims, faxes, images, chat transcripts, or call transcripts. Default for doc/comm work. Do NOT use BOT for these.
  - "DU"    = Document Understanding (the older product). Use ONLY when notes explicitly say "Document Understanding" or "DU". Otherwise use IXP for all doc/comm/form processing.
  Step description: 3-6 words, imperative. When a step orchestrates or routes work across agents/bots/humans, say Maestro. DO NOT invent steps, agents, bots, humans, or IXP steps that the notes don't mention. If the notes only describe 3 steps, return 3 steps.
- OUTCOMES: 1-5 outcome tiles. Each item: {value, label, source, note}. See transparency rules above. Value examples: "$558K", "90%", "9 min", or qualitative "Reduced" / "Faster" / "Fewer" / "Yes". Label examples: "revenue released", "cycle time", "of workflow automated". Empty array if the notes truly have no outcomes.
- ATTRIBUTABLE IMPACT (optional): list of directional metrics directly moved. Each item: {"direction": "up" | "down", "text": "metric name"}. 3-5 items. Empty list if none clearly inferable from the notes. Use "down" for reductions (cycle time, touches, backlogs) and "up" for improvements (yield, satisfaction, throughput). Only list metrics that the notes actually discuss or clearly imply — do NOT invent metrics the rep didn't mention.
- DOWNSTREAM IMPACT (optional): list of second-order effects (staff retention, NPS, compliance posture). Same format as attributable. Empty list if none inferable.
- BREADCRUMB: three items: [industry, function/department, use case name]
- COMPANY: customer company name
- THEME: "light" by default

Return ONLY a valid JSON object. No markdown fences, no preamble:
{
  "breadcrumb": ["industry", "function", "use case name"],
  "title": "string (short, 3-6 words)",
  "subtitle": "string",
  "company": "string",
  "problem_desc": "string",
  "problem_stats": [{"value": "string", "label": "string", "source": "stated|calculated|estimated|qualitative", "note": "string (required if calculated/estimated, else optional)"}],
  "solution_desc": "string",
  "capabilities": ["string", "..."],
  "steps": [{"role": "AGENT|BOT|HUMAN|IXP|DU", "description": "short step name"}],
  "outcomes": [{"value": "string", "label": "string", "source": "stated|calculated|estimated|qualitative", "note": "string (required if calculated/estimated, else optional)"}],
  "attributable": [{"direction": "up|down", "text": "metric name"}],
  "downstream": [{"direction": "up|down", "text": "metric name"}],
  "theme": "light"
}"""


SUGGESTIONS_PROMPT = """You are a seasoned CSM at UiPath. Your team is building 1-slide success stories for customers who have successfully deployed agentic automations into production. You are talking to an internal UiPath seller (CSM or AE), not the customer.

Your primary job: ensure the "Measured outcomes" section has quantitative stats that are most relevant and attention-grabbing to customer and executive readers. Review the slide data and recommend whether the current Measured outcomes are strong or whether better options exist given the context, and why. Make recommendations as concrete and metric-based as possible from the information provided.

Beyond Measured outcomes, you may also surface the highest-leverage gaps in the rest of the story — a missing problem stat, a fuzzy capability claim, a downstream effect worth naming. But Measured outcomes is the first thing you scrutinize.

Be brutally concise. Fewer is better. Skip anything obvious or already covered.

Return ONLY this JSON:
{
  "suggestions": [
    {
      "priority": "high" | "medium",
      "title": "short action (3-8 words, imperative)",
      "detail": "1 sentence on why this matters (max 20 words)",
      "tip": "where someone might look for this data (max 12 words, see list below)"
    }
  ]
}

For "tip", suggest plausible internal sources like: Salesforce opportunity record, Gainsight CSM notes, the AE, QBR/EBR deck, UiPath Insights, Automation Hub, the customer champion, internal Slack channels, product marketing. Pick what actually fits — don't force it. If a suggestion isn't about missing data (e.g. a reframe or angle), the tip can be an action instead (e.g. "Reframe the intro to lead with revenue unlocked").

Rules for what counts as a strong Measured outcome:
- A specific number paired with a meaningful business unit. "$558K revenue released" beats "90% automation". "Cycle time 6 hrs -> 9 min" beats "faster".
- Tied to outcomes a customer/executive cares about: revenue unlocked, capacity freed, decision quality, error/risk reduced, customer experience, compliance posture, time-to-cash.
- Three classes of outcome to look for: financial (revenue/cost/capacity), operational (cycle time, throughput, error rate, SLA attainment), experiential (CSAT/NPS, employee satisfaction, audit readiness).
- Watch for over-indexing on time savings only — that's the weakest framing for executives.
- If outcomes are vague ("improved efficiency", "faster processing"), suggest a concrete metric the rep should chase down.

Rules:
- MAX 4 items. Prefer 2-3.
- High priority only for the 1-2 most impactful.
- At least one suggestion should address Measured outcomes if there's any room to strengthen them.
- No filler. Every suggestion must be concretely actionable.
- Return ONLY the JSON. No markdown, no preamble."""


STEPS_PROMPT = """You are a UiPath process analyst. Given a description of an agentic automation process, break it into 3-9 discrete steps.

For each step, classify the actor role:
- "AGENT" = AI reasoning, classification, routing, decisions
- "BOT"   = deterministic RPA (data entry, API calls, portal polling, system updates)
- "HUMAN" = human in the loop (review, approve, sign-off)
- "IXP"   = documents and communications processing (extracting data from PDFs, emails, forms, invoices, claims, faxes, images). Default for doc/comm work. Use IXP — not BOT — for any step that parses unstructured documents.
- "DU"    = Document Understanding (older product). Use ONLY when the source notes explicitly say "Document Understanding" or "DU". Otherwise use IXP.

Step description: 3-6 words, imperative (e.g. "Pull & classify inbound denials", "Extract fields via IXP", "Route via Maestro"). Use Maestro for orchestration, routing, or handoff steps.

Return ONLY valid JSON:
{"steps": [{"role": "AGENT|BOT|HUMAN|IXP|DU", "description": "short step name"}]}"""


def _call_claude(system, user_text, max_tokens=2048, timeout=45):
    api_key = os.environ.get('ANTHROPIC_API_KEY', '').strip()
    if not api_key:
        raise ValueError('ANTHROPIC_API_KEY not configured on server.')
    resp = http_requests.post(
        'https://api.anthropic.com/v1/messages',
        headers={
            'x-api-key': api_key,
            'anthropic-version': '2023-06-01',
            'content-type': 'application/json',
        },
        json={
            'model': 'claude-haiku-4-5-20251001',
            'max_tokens': max_tokens,
            'system': system,
            'messages': [{'role': 'user', 'content': user_text}],
        },
        timeout=timeout,
    )
    if resp.status_code != 200:
        try:
            msg = resp.json().get('error', {}).get('message', resp.text)
        except Exception:
            msg = resp.text
        raise RuntimeError(f'Anthropic API error ({resp.status_code}): {msg}')
    raw = resp.json()['content'][0]['text'].strip()
    cleaned = raw.replace('```json', '').replace('```', '').strip()
    depth = 0
    start = cleaned.find('{')
    if start == -1:
        raise json.JSONDecodeError('No JSON object found', cleaned, 0)
    for i, ch in enumerate(cleaned[start:], start):
        if ch == '{': depth += 1
        elif ch == '}': depth -= 1
        if depth == 0:
            return json.loads(cleaned[start:i+1])
    return json.loads(cleaned)


@app.route('/')
def index():
    return send_from_directory('static', 'index.html')


@app.route('/<path:filename>')
def static_files(filename):
    return send_from_directory('static', filename)


_IXP_TRIGGERS = (
    # Product names / synonyms
    'doc understanding', 'document understanding', 'docunderstanding',
    'communications mining', 'communication mining', 'comms mining',
    'intelligent document', 'idp', 'ocr',
    # Document artifacts
    'pdf', 'invoice', 'claim', 'fax', 'form ', 'forms', 'questionnaire',
    'survey', 'application form', 'intake form', 'intake document',
    'contract', 'statement', 'remittance', 'rfq', 'rfp', 'po ', 'purchase order',
    'receipt', 'ack', 'asn', 'edi', 'attachment', 'packing slip', 'bill of lading',
    'letter', 'memo', 'scan', 'scanned', 'screenshot',
    # Communications artifacts
    ' email', 'emails', 'inbox', 'voicemail', 'call transcript', 'transcript',
    'chat message', 'sms', 'ticket text', 'support case', 'customer message',
    # Action verbs that signal IXP work
    'unstructured', 'extract field', 'extract fields', 'parse', 'read invoice',
    'read claim', 'read email', 'read pdf', 'read form', 'read document',
    'classify document', 'classify email', 'pull data from', 'pull fields',
    'ingest document', 'ingest email', 'ingest form', 'ingest invoice',
    'kyc', 'onboarding form',
)
_MAESTRO_TRIGGERS = (
    # Explicit
    'maestro', 'orchestrat', 'process orchestrator', 'workflow engine',
    'workload registration', 'workflow registration', 'process controller',
    # Coordination patterns
    'route', 'routing', 'dispatch', 'handoff', 'hand off', 'hand-off',
    'handover', 'hand over', 'hand-over', 'sla manage', 'sla-manage',
    'escalat', 'queue manage', 'priorit', 'workload', 'coordinat',
)


def _enforce_capability_rules(parsed, source_text):
    """Force IXP/DU and Maestro into the output if the source text or steps imply them.

    DU vs IXP routing: explicit "Document Understanding" / "DU" in source notes
    means use DU (older product). All other doc/comm/form processing -> IXP.
    """
    if not isinstance(parsed, dict):
        return parsed
    src = (source_text or '').lower()
    caps = list(parsed.get('capabilities') or [])

    # Detect whether the user EXPLICITLY named DU. Only these literal phrases count.
    du_explicit = any(t in src for t in (
        'document understanding', 'doc understanding', 'docunderstanding',
        ' du ', ' du.', ' du,', ' du)', ' du:', '(du)',
    ))

    # Map any "Document Understanding" / "Communications Mining" capabilities the
    # AI emitted into the canonical name (DU if explicit, else IXP). Communications
    # Mining always collapses to IXP.
    canonical_caps = []
    seen = set()
    for c in caps:
        cl = str(c).lower().strip()
        if cl in ('doc understanding', 'document understanding', 'docunderstanding'):
            target = 'DU' if du_explicit else 'IXP'
        elif cl in ('communications mining', 'communication mining', 'comms mining'):
            target = 'IXP'
        else:
            target = c
        key = str(target).lower().strip()
        if key not in seen:
            canonical_caps.append(target); seen.add(key)
    caps = canonical_caps
    caps_lower = [str(c).lower() for c in caps]

    # Doc/comm/form processing enforcement
    doc_in_text = any(t in src for t in _IXP_TRIGGERS) or du_explicit
    has_du = 'du' in caps_lower
    has_ixp = any('ixp' == c.strip() or 'ixp' in c.split() for c in caps_lower)
    if doc_in_text and not (has_du or has_ixp):
        caps.append('DU' if du_explicit else 'IXP')
        caps_lower.append('du' if du_explicit else 'ixp')

    # Maestro enforcement — required if triggers in text OR steps span 3+ different roles
    maestro_in_text = any(t in src for t in _MAESTRO_TRIGGERS)
    maestro_in_caps = any('maestro' in c for c in caps_lower)
    steps = parsed.get('steps') or []
    role_set = {str(s.get('role','')).upper() for s in steps if isinstance(s, dict)}
    multi_role = len(role_set & {'AGENT','BOT','HUMAN','IXP'}) >= 2 and len(steps) >= 3
    if (maestro_in_text or multi_role) and not maestro_in_caps:
        caps.append('Maestro')

    parsed['capabilities'] = caps

    # Step-level: rewrite ANY step that looks like document/communication/form
    # intake as IXP, regardless of original role (BOT or AGENT).
    doc_action_words = (
        'extract', 'parse', 'read', 'classify', 'intake', 'ingest', 'triage',
        'retrieve', 'fetch', 'pull', 'process', 'capture', 'index', 'ocr',
        'analyze', 'analyse', 'review', 'scan', 'recognize', 'recognise',
        'identify field', 'identify fields',
    )
    doc_object_words = (
        # Documents
        'document', 'doc ', 'pdf', 'invoice', 'claim', 'form', 'fax',
        'transcript', 'ack', 'asn', 'edi', 'attachment', 'po ', 'purchase order',
        'receipt', 'statement', 'contract', 'rfq', 'rfp', 'remittance',
        'application', 'letter', 'memo', 'questionnaire', 'survey', 'submission',
        'packing slip', 'bill of lading', 'image', 'scanned',
        # Communications
        'email', 'inbox', 'voicemail', 'call note', 'chat', 'sms', 'ticket',
        'support case', 'customer message', 'communication', 'comms ',
        # Other
        'ixp', 'kyc',
    )
    for s in steps:
        if not isinstance(s, dict):
            continue
        role = str(s.get('role','')).upper()
        if role in ('IXP', 'DU'):
            continue
        desc = str(s.get('description','')).lower()
        # If step describes acting on a doc/comm artifact, route to DU (when
        # explicit) or IXP (default).
        if any(t in desc for t in doc_object_words) and any(w in desc for w in doc_action_words):
            s['role'] = 'DU' if du_explicit else 'IXP'

    return parsed


@app.route('/extract', methods=['POST'])
def extract():
    """Takes raw text notes, returns structured JSON for the agentic template."""
    try:
        body = request.get_json(force=True)
        text = (body.get('text') or '').strip()
        if not text:
            return jsonify(error='No input text provided.'), 400
        parsed = _call_claude(EXTRACTION_PROMPT, f'Rep notes:\n\n{text}')
        parsed = _enforce_capability_rules(parsed, text)
        return jsonify(parsed)
    except json.JSONDecodeError as e:
        return jsonify(error=f'Could not parse AI response as JSON: {e}'), 500
    except http_requests.Timeout:
        return jsonify(error='AI request timed out. Try again.'), 504
    except Exception as e:
        traceback.print_exc()
        return jsonify(error=f'Extraction failed: {str(e)}'), 500


@app.route('/extract-steps', methods=['POST'])
def extract_steps():
    """Takes a process description, returns just the steps list."""
    try:
        body = request.get_json(force=True)
        text = (body.get('text') or '').strip()
        if not text:
            return jsonify(error='No process description provided.'), 400
        parsed = _call_claude(STEPS_PROMPT, f'Process description:\n\n{text}',
                              max_tokens=1024, timeout=30)
        return jsonify(parsed)
    except json.JSONDecodeError as e:
        return jsonify(error=f'Could not parse AI response as JSON: {e}'), 500
    except http_requests.Timeout:
        return jsonify(error='AI request timed out. Try again.'), 504
    except Exception as e:
        traceback.print_exc()
        return jsonify(error=f'Step extraction failed: {str(e)}'), 500


@app.route('/suggest', methods=['POST'])
def suggest():
    """Takes current story JSON, returns AI suggestions for strengthening it."""
    try:
        body = request.get_json(force=True) or {}
        story_json = json.dumps(body, ensure_ascii=False)
        parsed = _call_claude(SUGGESTIONS_PROMPT,
                              f'Current story data:\n\n{story_json}',
                              max_tokens=1024, timeout=30)
        return jsonify(parsed)
    except json.JSONDecodeError as e:
        return jsonify(error=f'Could not parse AI response as JSON: {e}'), 500
    except http_requests.Timeout:
        return jsonify(error='AI request timed out. Try again.'), 504
    except Exception as e:
        traceback.print_exc()
        return jsonify(error=f'Suggestions failed: {str(e)}'), 500


def _post_to_roi_webhook(body):
    """If ROI_WEBHOOK_URL env var is set, POST a flattened summary of this generation.
    Silent no-op when unset. All errors swallowed — never block a download."""
    url = os.environ.get('ROI_WEBHOOK_URL', '').strip()
    if not url:
        return
    try:
        bc = body.get('breadcrumb') or []
        outcomes = body.get('outcomes') or []
        capabilities = body.get('capabilities') or []
        steps = body.get('steps') or []
        payload = {
            'timestamp_utc': __import__('datetime').datetime.utcnow().isoformat() + 'Z',
            'company': body.get('company', ''),
            'title': body.get('title', ''),
            'industry': bc[0] if len(bc) > 0 else '',
            'function': bc[1] if len(bc) > 1 else '',
            'use_case': bc[2] if len(bc) > 2 else '',
            'classification': body.get('classification', ''),
            'account_team': body.get('account_team', ''),
            'capabilities': ', '.join(str(c) for c in capabilities),
            'capability_count': len(capabilities),
            'step_count': len(steps),
            'outcomes': '; '.join(
                f"{(o.get('value','') or '').strip()} {(o.get('label','') or '').strip()}"
                for o in outcomes if isinstance(o, dict)
            ),
            'outcome_count': len(outcomes),
            'has_maestro': any('maestro' in str(c).lower() for c in capabilities),
            'has_ixp': any('ixp' == str(c).lower().strip() or 'ixp' in str(c).lower().split()
                           for c in capabilities),
        }
        http_requests.post(url, json=payload, timeout=5)
    except Exception:
        pass


@app.route('/generate', methods=['POST'])
def generate():
    """Takes structured story JSON, returns the built .pptx as a download."""
    try:
        body = request.get_json(force=True) or {}
        result = build_pptx(body)
        pptx_bytes = result[0] if isinstance(result, tuple) else result
        company = (body.get('company') or 'win_story').strip()
        safe = ''.join(c if c.isalnum() or c in ('_', '-') else '_' for c in company)
        filename = f'{safe}_win_story.pptx'
        # Fire-and-forget tracking
        _post_to_roi_webhook(body)
        return Response(
            pptx_bytes,
            mimetype='application/vnd.openxmlformats-officedocument.presentationml.presentation',
            headers={'Content-Disposition': f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        traceback.print_exc()
        return jsonify(error=f'Generate failed: {str(e)}'), 500


def _extract_text_from_upload(filename, blob):
    """Best-effort text extraction from .pdf, .docx, .txt, .md."""
    name = (filename or '').lower()
    if name.endswith('.pdf'):
        try:
            import io as _io
            from pypdf import PdfReader
            reader = PdfReader(_io.BytesIO(blob))
            return '\n\n'.join((p.extract_text() or '') for p in reader.pages).strip()
        except Exception as e:
            return f"[PDF parse error: {e}]"
    if name.endswith('.docx'):
        try:
            import io as _io
            import docx
            d = docx.Document(_io.BytesIO(blob))
            parts = [p.text for p in d.paragraphs if p.text.strip()]
            for tbl in d.tables:
                for row in tbl.rows:
                    parts.append(' | '.join(c.text for c in row.cells))
            return '\n'.join(parts).strip()
        except Exception as e:
            return f"[DOCX parse error: {e}]"
    if name.endswith('.txt') or name.endswith('.md'):
        try:
            return blob.decode('utf-8', errors='replace').strip()
        except Exception as e:
            return f"[Text decode error: {e}]"
    return f"[Unsupported file type: {filename}]"


@app.route('/parse-docs', methods=['POST'])
def parse_docs():
    """Accept multiple uploaded files, return concatenated extracted text per file."""
    try:
        files = request.files.getlist('files')
        if not files:
            return jsonify(error='No files uploaded.'), 400
        results = []
        MAX_BYTES = 10 * 1024 * 1024  # 10 MB per file
        for f in files:
            blob = f.read(MAX_BYTES + 1)
            if len(blob) > MAX_BYTES:
                results.append({'name': f.filename, 'text': f'[File too large: {f.filename}, max 10MB]'})
                continue
            text = _extract_text_from_upload(f.filename, blob)
            if len(text) > 20000:
                text = text[:20000] + '\n[...truncated...]'
            results.append({'name': f.filename, 'text': text})
        return jsonify(files=results)
    except Exception as e:
        traceback.print_exc()
        return jsonify(error=f'Parse failed: {str(e)}'), 500


@app.route('/health')
def health():
    return jsonify(status='ok', has_api_key=bool(os.environ.get('ANTHROPIC_API_KEY', '').strip()))


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
